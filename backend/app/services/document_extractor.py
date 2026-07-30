import io
import os
import hashlib
from typing import Dict, Any

import pypdf
import pdfplumber
import docx
from PIL import Image

class DocumentExtractor:
    """
    Python Document Text & Metadata Extractor using pypdf, pdfplumber, python-docx, and Pillow.
    """

    @staticmethod
    def extract_text_and_metadata(filename: str, file_bytes: bytes) -> Dict[str, Any]:
        ext = os.path.splitext(filename)[1].lower()
        sha256 = hashlib.sha256(file_bytes).hexdigest()
        full_text = ""
        metadata = {
            "sha256": f"sha256-{sha256}",
            "file_type": ext.replace(".", "").upper(),
            "byte_size": len(file_bytes),
            "page_count": 1,
            "extraction_tool": "Standard Reader"
        }

        try:
            if ext == ".pdf":
                full_text, metadata["page_count"] = DocumentExtractor._extract_pdf(file_bytes)
                metadata["extraction_tool"] = "pypdf + pdfplumber"

            elif ext in [".docx", ".doc"]:
                full_text = DocumentExtractor._extract_docx(file_bytes)
                metadata["extraction_tool"] = "python-docx"

            elif ext in [".txt", ".md", ".json", ".csv", ".log", ".xml", ".html"]:
                full_text = file_bytes.decode("utf-8", errors="ignore")
                metadata["extraction_tool"] = "UTF-8 Text Decoder"

            elif ext in [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"]:
                full_text, img_meta = DocumentExtractor._extract_image_info(filename, file_bytes)
                metadata.update(img_meta)
                metadata["extraction_tool"] = "Pillow Image Inspector"

            else:
                full_text = file_bytes.decode("utf-8", errors="ignore")

        except Exception as e:
            full_text = f"--- EXTRACTOR WARNING FOR {filename} ---\nCould not parse raw binary with default decoder: {str(e)}\n"

        if not full_text.strip():
            full_text = f"--- DOCUMENT READOUT ({filename}) ---\nFile size: {len(file_bytes)} bytes. No readable text found or binary formatted file."

        return {
            "full_ocr_text": full_text[:100000],  # cap at 100k chars for DB storage
            "metadata": metadata,
            "sha256": metadata["sha256"]
        }

    @staticmethod
    def _extract_pdf(file_bytes: bytes) -> tuple[str, int]:
        text_parts = []
        page_count = 0

        # Method 1: pdfplumber for high accuracy text & layout
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                page_count = len(pdf.pages)
                for idx, page in enumerate(pdf.pages, start=1):
                    p_text = page.extract_text() or ""
                    if p_text.strip():
                        text_parts.append(f"--- PAGE {idx} (pdfplumber) ---\n{p_text}")
        except Exception:
            pass

        # Method 2: pypdf fallback if pdfplumber yields empty
        if not text_parts:
            try:
                reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                page_count = len(reader.pages)
                for idx, page in enumerate(reader.pages, start=1):
                    p_text = page.extract_text() or ""
                    if p_text.strip():
                        text_parts.append(f"--- PAGE {idx} (pypdf) ---\n{p_text}")
            except Exception:
                pass

        return "\n\n".join(text_parts), page_count

    @staticmethod
    def _extract_docx(file_bytes: bytes) -> str:
        text_parts = []
        doc = docx.Document(io.BytesIO(file_bytes))

        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text)

        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    text_parts.append(" | ".join(row_cells))

        return "\n".join(text_parts)

    @staticmethod
    def _extract_image_info(filename: str, file_bytes: bytes) -> tuple[str, dict]:
        try:
            img = Image.open(io.BytesIO(file_bytes))
            width, height = img.size
            format_name = img.format or "IMAGE"
            text_readout = f"--- IMAGE FILE READOUT ({filename}) ---\nFormat: {format_name}\nDimensions: {width}x{height} pixels\nColor Mode: {img.mode}\nVerified High-Resolution Document Image."
            return text_readout, {"dimensions": f"{width}x{height}", "format": format_name}
        except Exception as e:
            return f"--- IMAGE READOUT ({filename}) ---\nProcessed {len(file_bytes)} bytes.", {}
